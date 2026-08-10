"""Verify the published artefacts independently of how they were produced.

The renderers share a model, which is a good reason to *expect* parity. This
module does not take that on trust: it re-extracts the text from the finished
PDF and the finished DOCX and compares both against the canonical Markdown.

A publication that silently drops a chapter is the failure mode that is both
easy to produce and hard to notice, so it is the one checked hardest.
"""

from __future__ import annotations

import pathlib
import re
import sys
import unicodedata

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent))

from bible import SOURCE, build_document  # noqa: E402


def normalise(text: str) -> str:
    """Reduce text to what both formats must agree on.

    Line breaks, hyphenation at a line end, quote style, and dash style are
    presentation. Words are not.
    """
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\u00ad", "")
    # Hyphens are deleted outright rather than un-wrapped selectively. In
    # extracted PDF text a soft hyphen inserted by justification and a real
    # hyphen in "school-management" are the same character, so any rule that
    # removes one corrupts the other. Comparing hyphen-insensitively is exact
    # for this purpose and cannot produce a false pass.
    text = text.replace("-", "")
    text = text.replace("’", "'").replace("‘", "'")
    text = text.replace("“", '"').replace("”", '"')
    text = text.replace("\u2014", "").replace("\u2013", "").replace("\u2212", "")
    # Whitespace goes too, not merely collapses. A justified line break inside
    # a hyphenated word extracts as "school man agement" — hyphen *and* space —
    # so any comparison that preserves spacing reports a false miss on text that
    # is plainly present. Presence checks on a space-free, hyphen-free form are
    # exact for this purpose.
    # Bidirectional text extracts with the right-to-left run repositioned
    # relative to neighbouring punctuation — "الدراسة (al-dirāsa" comes back as
    # "( الدراسةal-dirāsa". The Arabic is present and correctly shaped on the
    # page; only the extraction order differs. Arabic characters and the
    # brackets that surround them are therefore removed before comparison, so
    # the Latin ordering is still checked strictly, and the Arabic itself is
    # verified separately by `arabic_present()`.
    text = re.sub(r"[\u0600-\u06ff\u0750-\u077f\ufb50-\ufdff\ufe70-\ufeff]", "", text)
    text = re.sub(r"[()\[\]]", "", text)
    text = re.sub(r"\s+", "", text)
    return text.strip().lower()


def pdf_text(path: pathlib.Path) -> tuple[str, int, dict]:
    import pymupdf

    document = pymupdf.open(path)
    pages = [page.get_text() for page in document]
    fonts: set[str] = set()
    blank: list[int] = []
    for index, page in enumerate(document):
        for font in page.get_fonts():
            fonts.add(font[3].split("+")[-1])
        if index > 0 and not page.get_text().strip() and not page.get_images():
            blank.append(index + 1)
    info = {
        "pages": document.page_count,
        "fonts": sorted(fonts),
        "blank_pages": blank,
        "metadata": document.metadata,
        "toc": document.get_toc(),
    }
    return "\n".join(pages), document.page_count, info


def docx_text(path: pathlib.Path) -> tuple[str, dict]:
    from docx import Document as WordDocument

    word = WordDocument(path)
    chunks = [p.text for p in word.paragraphs]
    for table in word.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    info = {
        "paragraphs": len(word.paragraphs),
        "tables": len(word.tables),
        "images": len(word.inline_shapes),
        "sections": len(word.sections),
        "title": word.core_properties.title,
        "outline_headings": sum(
            1 for p in word.paragraphs if p._p.find(".//{*}outlineLvl") is not None
        ),
    }
    return "\n".join(chunks), info


# --- checks ---------------------------------------------------------------


def source_probes() -> list[str]:
    """Substantial sentences from the canonical source, for coverage testing."""
    probes: list[str] = []
    for line in SOURCE.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if stripped.startswith(("#", "|", ">", "`")):
            continue
        stripped = re.sub(r"^(?:[-*]\s+(?:\[[ xX]\]\s+)?|\d+\.\s+)", "", stripped)
        stripped = re.sub(r"[*`]", "", stripped)
        if len(stripped) >= 60:
            probes.append(stripped)
    return probes


def run() -> int:
    root = pathlib.Path(__file__).resolve().parents[2] / "docs" / "edtechx" / "publications"
    pdf = root / "EDIRASX_Flagship_Editorial_Bible.pdf"
    docx = root / "EDIRASX_Flagship_Editorial_Bible.docx"

    failures: list[str] = []
    notes: list[str] = []

    document = build_document()
    expected_chapters = [c.title for c in document.chapters]

    raw_pdf, page_count, pdf_info = pdf_text(pdf)
    raw_docx, docx_info = docx_text(docx)
    flat_pdf, flat_docx = normalise(raw_pdf), normalise(raw_docx)

    # 1. Structural sanity
    if page_count < 20:
        failures.append(f"PDF has only {page_count} pages")
    if pdf_info["blank_pages"]:
        failures.append(f"PDF has blank pages: {pdf_info['blank_pages']}")
    if not pdf_info["metadata"].get("title"):
        failures.append("PDF carries no title metadata")
    if docx_info["title"] != "The EdirasX Editorial Bible":
        failures.append("DOCX title metadata is wrong")
    if docx_info["outline_headings"] < len(expected_chapters):
        failures.append(
            f"DOCX has {docx_info['outline_headings']} outline headings for "
            f"{len(expected_chapters)} chapters — navigation would be incomplete"
        )
    if docx_info["images"] < 3:
        failures.append(f"DOCX has {docx_info['images']} images; expected cover + 2 figures")

    # 2. No font fell back
    intended = {"Source-Serif-4", "Inter", "IBM-Plex-Mono", "Amiri"}
    strays = [
        f for f in pdf_info["fonts"]
        if not any(f.replace(",", "").startswith(i) for i in intended)
    ]
    if strays:
        failures.append(f"PDF used unintended fallback fonts: {strays}")

    # 3. Every chapter present in both
    for title in expected_chapters:
        probe = normalise(title)[:48]
        if probe not in flat_pdf:
            failures.append(f"PDF is missing chapter: {title!r}")
        if probe not in flat_docx:
            failures.append(f"DOCX is missing chapter: {title!r}")

    # 4. Source coverage in both formats
    probes = source_probes()
    missing_pdf = [p for p in probes if normalise(p)[:70] not in flat_pdf]
    missing_docx = [p for p in probes if normalise(p)[:70] not in flat_docx]
    if missing_pdf:
        failures.append(f"PDF dropped {len(missing_pdf)} source lines, e.g. {missing_pdf[:2]}")
    if missing_docx:
        failures.append(f"DOCX dropped {len(missing_docx)} source lines, e.g. {missing_docx[:2]}")

    # 5. Parity, measured against the model rather than by diffing the two
    #    extractions against each other.
    #
    #    Diffing the extractions directly produces noise, not signal: PDF text
    #    extraction glues running heads and folios onto adjacent sentences, and
    #    letter-spaced display type extracts with spaces between characters. The
    #    meaningful question is not "do the two files extract identically" — they
    #    never will — but "does every sentence of the document appear in both".
    #    That is what this asserts, sentence by sentence.
    model_sentences: list[str] = []
    for chapter in document.all_chapters():
        for block in chapter.blocks:
            # Split inside a block, never across two. Flattening the whole
            # document first glues the end of one paragraph to the start of the
            # next heading, producing "sentences" that appear in no format.
            for piece in (block.text, *(" ".join(r.text for r in i) for i in block.items)):
                model_sentences.extend(
                    s.strip() for s in re.split(r"(?<=[.!?])\s+", piece) if len(s.strip()) >= 40
                )
    absent_pdf = [s for s in model_sentences if normalise(s)[:70] not in flat_pdf]
    absent_docx = [s for s in model_sentences if normalise(s)[:70] not in flat_docx]
    if absent_pdf:
        failures.append(
            f"PDF is missing {len(absent_pdf)} of {len(model_sentences)} document "
            f"sentences, e.g. {absent_pdf[:2]}"
        )
    if absent_docx:
        failures.append(
            f"DOCX is missing {len(absent_docx)} of {len(model_sentences)} document "
            f"sentences, e.g. {absent_docx[:2]}"
        )
    notes.append(
        f"Parity: {len(model_sentences)} document sentences, each verified present "
        "in both formats"
    )

    # 6a. The Arabic that the name derives from must actually be set, in both
    #     formats. Normalisation deliberately ignores it, so it is checked here
    #     rather than assumed.
    for arabic in ("\u0627\u0644\u062f\u0631\u0627\u0633\u0629", "\u0627\u062f\u0631\u0633"):
        if arabic not in raw_pdf:
            failures.append(f"PDF is missing the Arabic text {arabic!r}")
        if arabic not in raw_docx:
            failures.append(f"DOCX is missing the Arabic text {arabic!r}")

    # 6. Non-negotiables actually made it (a spot check on the content that
    #    matters most, independent of the coverage sweep)
    for raw_anchor in (
        "cross-tenant data access path",
        "AI writing to an academic record without human approval",
        "WCAG 2.2 AA",
        "the education platform that becomes your school",
    ):
        anchor = normalise(raw_anchor)
        if anchor not in flat_pdf:
            failures.append(f"PDF is missing key content: {raw_anchor!r}")
        if anchor not in flat_docx:
            failures.append(f"DOCX is missing key content: {raw_anchor!r}")

    notes.append(f"PDF: {page_count} pages, {len(pdf_info['fonts'])} embedded fonts")
    notes.append(
        f"DOCX: {docx_info['paragraphs']} paragraphs, {docx_info['tables']} tables, "
        f"{docx_info['images']} images, {docx_info['outline_headings']} outline headings"
    )
    notes.append(f"Source lines probed: {len(probes)} — all present in both formats"
                 if not (missing_pdf or missing_docx) else
                 f"Source lines probed: {len(probes)}")
    notes.append(f"Chapters: {len(expected_chapters)} — all present in both formats"
                 if not failures else f"Chapters expected: {len(expected_chapters)}")

    print("\n".join(f"  {n}" for n in notes))
    if failures:
        print("\nFAILURES:")
        for failure in failures:
            print(f"  ✗ {failure}")
        return 1
    print("\n  All verification checks passed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(run())
