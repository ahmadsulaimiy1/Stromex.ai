"""Render the document model to a premium Word document.

Same model as the PDF, so the two carry identical text by construction rather
than by coincidence. What differs is only what must: Word gets real outline
levels so the Navigation pane works, a field-based table of contents, and a
rasterised cover and figures, because Word has no dependable SVG support.
"""

from __future__ import annotations

import io
import pathlib

from docx import Document as WordDocument
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor

from diagrams import DIAGRAMS
from model import Block, BlockKind, Chapter, Document, Run, typographic

ASSETS = pathlib.Path(__file__).resolve().parent / "assets"

SERIF = "Source Serif 4"
SANS = "Inter"
MONO = "IBM Plex Mono"

INK = RGBColor(0x16, 0x19, 0x1C)
INK_2 = RGBColor(0x40, 0x47, 0x4E)
INK_3 = RGBColor(0x6B, 0x71, 0x78)
ACCENT = RGBColor(0x16, 0x32, 0x4F)
ACCENT_2 = RGBColor(0x2E, 0x6E, 0xA8)
VIOLET = RGBColor(0x7A, 0x3F, 0xD6)


# --- low-level helpers ----------------------------------------------------


def _set_font(element, name: str) -> None:
    """Pin a font on a run or on a style's run-properties.

    Word only honours a font it is told about on every script axis; setting
    `ascii` alone leaves the complex-script and East-Asian axes on the theme
    font, which is how a document ends up half in the intended face.
    """
    rPr = element.get_or_add_rPr() if hasattr(element, "get_or_add_rPr") else element
    fonts = rPr.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "cs", "eastAsia"):
        fonts.set(qn(f"w:{attribute}"), name)


def _shade(element, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    element.append(shading)


def _border(paragraph, edge: str, colour: str, size: int = 18) -> None:
    borders = OxmlElement("w:pBdr")
    element = OxmlElement(f"w:{edge}")
    element.set(qn("w:val"), "single")
    element.set(qn("w:sz"), str(size))
    element.set(qn("w:space"), "8")
    element.set(qn("w:color"), colour)
    borders.append(element)
    paragraph._p.get_or_add_pPr().append(borders)


def _field(paragraph, instruction: str, placeholder: str = "") -> None:
    """Insert a Word field (used for the contents and for page numbers).

    Each part of a field goes in its own run. Packing begin/instrText/separate/
    end into a single run parses fine with a permissive library and then fails
    to open in Word and LibreOffice, which is the worst of both worlds: no
    error at build time, an unusable file at delivery.
    """

    def _run_with(child):
        run = paragraph.add_run()
        run._r.append(child)
        return run

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    _run_with(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    _run_with(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    _run_with(separate)

    text = OxmlElement("w:t")
    text.text = placeholder
    _run_with(text)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    _run_with(end)


def _outline_level(paragraph, level: int) -> None:
    """Give a paragraph an outline level so Word's Navigation pane sees it."""
    pPr = paragraph._p.get_or_add_pPr()
    element = OxmlElement("w:outlineLvl")
    element.set(qn("w:val"), str(level))
    pPr.append(element)


def _bookmark(paragraph, name: str, index: int) -> None:
    """Anchor a bookmark on a paragraph.

    The start marker goes *after* `w:pPr`, never before it: paragraph
    properties must be the first child of `w:p`, and getting that wrong
    produces a file that opens nowhere.
    """
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(index))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(index))

    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(start)
    else:
        paragraph._p.insert(0, start)
    paragraph._p.append(end)


# --- text -----------------------------------------------------------------


def _add_runs(paragraph, runs: tuple[Run, ...], *, size: Pt, colour=INK, font=SERIF) -> None:
    for run in runs:
        added = paragraph.add_run(run.text)
        added.bold = run.bold
        added.italic = run.italic
        added.font.size = Pt(size.pt - 0.6) if run.code else size
        added.font.color.rgb = colour
        _set_font(added._r, MONO if run.code else font)


def _paragraph(container, *, space_after: float = 6, space_before: float = 0, line: float = 1.25):
    paragraph = container.add_paragraph()
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(space_after)
    fmt.space_before = Pt(space_before)
    fmt.line_spacing = line
    return paragraph


# --- blocks ---------------------------------------------------------------


def _render_block(doc, block: Block) -> None:
    kind = block.kind

    if kind is BlockKind.heading:
        sizes = {2: 14, 3: 11, 4: 10}
        size = sizes.get(min(max(block.level, 2), 4), 10)
        paragraph = _paragraph(doc, space_before=14, space_after=5)
        if block.number:
            number = paragraph.add_run(f"{block.number}  ")
            number.font.size = Pt(size - 1)
            number.font.color.rgb = INK_3
            _set_font(number._r, SANS)
        colour = INK if block.level == 2 else ACCENT
        _add_runs(
            paragraph,
            block.runs,
            size=Pt(size),
            colour=colour,
            font=SERIF if block.level == 2 else SANS,
        )
        for run in paragraph.runs:
            run.bold = True
        _outline_level(paragraph, min(block.level, 4))
        return

    if kind is BlockKind.paragraph:
        paragraph = _paragraph(doc)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_runs(paragraph, block.runs, size=Pt(10.5))
        return

    if kind is BlockKind.quote:
        paragraph = _paragraph(doc, space_before=10, space_after=12)
        paragraph.paragraph_format.left_indent = Inches(0.28)
        _border(paragraph, "left", "7A3FD6", size=24)
        _add_runs(paragraph, block.runs, size=Pt(13), colour=ACCENT)
        for run in paragraph.runs:
            run.bold = True
        return

    if kind in (BlockKind.bullets, BlockKind.numbers, BlockKind.checklist):
        for index, item in enumerate(block.items, start=1):
            paragraph = _paragraph(doc, space_after=3)
            paragraph.paragraph_format.left_indent = Inches(0.32)
            paragraph.paragraph_format.first_line_indent = Inches(-0.32)
            if kind is BlockKind.numbers:
                marker = paragraph.add_run(f"{index} ")
                marker.font.color.rgb = ACCENT_2
                marker.bold = True
                _set_font(marker._r, SANS)
                marker.font.size = Pt(9)
            elif kind is BlockKind.checklist:
                marker = paragraph.add_run("☐ ")
                marker.font.color.rgb = INK_3
                marker.font.size = Pt(11)
            else:
                marker = paragraph.add_run("— ")
                marker.font.color.rgb = ACCENT_2
                marker.font.size = Pt(10.5)
            _add_runs(paragraph, item, size=Pt(10.5))
        return

    if kind is BlockKind.table:
        table = doc.add_table(rows=1, cols=len(block.header))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = True
        for cell, runs in zip(table.rows[0].cells, block.header, strict=False):
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            _add_runs(paragraph, runs, size=Pt(8), colour=INK_3, font=SANS)
            for run in paragraph.runs:
                run.bold = True
            _shade(cell._tc.get_or_add_tcPr(), "F5F5F1")
        for row_runs in block.rows:
            cells = table.add_row().cells
            for cell, runs in zip(cells, row_runs, strict=False):
                cell.text = ""
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_after = Pt(2)
                _add_runs(paragraph, runs, size=Pt(9), font=SANS)
        _paragraph(doc, space_after=8)
        return

    if kind is BlockKind.callout:
        if block.label:
            label = _paragraph(doc, space_before=10, space_after=2)
            label.paragraph_format.left_indent = Inches(0.16)
            run = label.add_run(block.label.upper())
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = ACCENT_2
            _set_font(run._r, SANS)
            _border(label, "left", "2E6EA8", size=24)
        body = _paragraph(doc, space_after=12)
        body.paragraph_format.left_indent = Inches(0.16)
        _border(body, "left", "2E6EA8", size=24)
        _add_runs(body, block.runs, size=Pt(10.5))
        return

    if kind is BlockKind.contrast:
        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        headers = ("Prestige is not", "Prestige is")
        for index, title in enumerate(headers):
            cell = table.rows[0].cells[index]
            cell.text = ""
            run = cell.paragraphs[0].add_run(title.upper())
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x9B, 0x22, 0x26) if index == 0 else RGBColor(0x1C, 0x6B, 0x4A)
            _set_font(run._r, SANS)
            _shade(cell._tc.get_or_add_tcPr(), "F5F5F1")
        left = [" ".join(r.text for r in item) for item in block.items]
        right = [" ".join(r.text for r in row[0]) for row in block.rows]
        for index, values in enumerate((left, right)):
            cell = table.rows[1].cells[index]
            cell.text = ""
            for position, value in enumerate(values):
                paragraph = cell.paragraphs[0] if position == 0 else cell.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(2)
                run = paragraph.add_run(f"·  {value}")
                run.font.size = Pt(9.5)
                _set_font(run._r, SERIF)
        _paragraph(doc, space_after=8)
        return

    if kind is BlockKind.diagram:
        svg = DIAGRAMS.get(block.name)
        if svg:
            import cairosvg

            png = cairosvg.svg2png(bytestring=svg.encode(), output_width=1800)
            doc.add_picture(io.BytesIO(png), width=Inches(6.3))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if block.label:
            caption = _paragraph(doc, space_before=4, space_after=12)
            run = caption.add_run(typographic(block.label))
            run.font.size = Pt(8)
            run.font.color.rgb = INK_3
            _set_font(run._r, SANS)
        return

    if kind is BlockKind.rule:
        paragraph = _paragraph(doc, space_before=8, space_after=8)
        _border(paragraph, "bottom", "D9D9D3", size=6)
        return


# --- document assembly ----------------------------------------------------


def _configure_styles(word) -> None:
    normal = word.styles["Normal"]
    normal.font.name = SERIF
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    _set_font(normal.element.get_or_add_rPr(), SERIF)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25


def _cover(word, pdf_path: pathlib.Path) -> None:
    """Use the PDF's own cover, rasterised.

    Word cannot reproduce a full-bleed dark page reliably across versions, and
    a cover that renders differently in each format would be the most visible
    possible parity failure. Rendering the identical artwork removes the
    question.
    """
    import pymupdf

    document = pymupdf.open(pdf_path)
    pixmap = document[0].get_pixmap(dpi=200)
    image = io.BytesIO(pixmap.tobytes("png"))

    section = word.sections[0]
    section.top_margin = section.bottom_margin = Emu(0)
    section.left_margin = section.right_margin = Emu(0)
    paragraph = word.paragraphs[0] if word.paragraphs else word.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(image, width=section.page_width)


def _body_section(word):
    section = word.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.05)
    section.right_margin = Inches(0.9)

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _field(paragraph, " PAGE ", "1")
    for run in paragraph.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = INK_3
        _set_font(run._r, SANS)

    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("EdirasX Editorial Bible")
    run.font.size = Pt(7.5)
    run.font.color.rgb = INK_3
    _set_font(run._r, SANS)
    return section


def _chapter_heading(word, chapter: Chapter, eyebrow: str, index: int) -> None:
    word.add_page_break()
    if eyebrow:
        paragraph = _paragraph(word, space_after=4)
        run = paragraph.add_run(eyebrow.upper())
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = ACCENT_2
        _set_font(run._r, SANS)
    title = _paragraph(word, space_after=10)
    run = title.add_run(typographic(chapter.title))
    run.bold = True
    run.font.size = Pt(21)
    run.font.color.rgb = INK
    _set_font(run._r, SERIF)
    _outline_level(title, 0)
    _bookmark(title, f"ch{index}", index)
    _border(title, "bottom", "7A3FD6", size=12)


def render(document: Document, output: pathlib.Path) -> pathlib.Path:
    pdf_path = output.with_suffix(".pdf")
    word = WordDocument()
    _configure_styles(word)

    core = word.core_properties
    core.title = "The EdirasX Editorial Bible"
    core.subject = document.subtitle
    core.author = "EdirasX"
    core.category = "Foundational document"
    core.comments = (
        "Official publication edition. Generated from "
        "docs/edtechx/EDTECHX_EDITORIAL_BIBLE.md, which remains canonical."
    )
    core.keywords = "EdirasX; Editorial Bible; education platform; constitution"

    _cover(word, pdf_path)
    _body_section(word)

    # Document control
    heading = _paragraph(word, space_after=10)
    run = heading.add_run("DOCUMENT CONTROL")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = INK_3
    _set_font(run._r, SANS)
    _outline_level(heading, 0)

    control = word.add_table(rows=0, cols=2)
    control.style = "Table Grid"
    for key, value in (
        ("Document", "The EdirasX Editorial Bible"),
        ("Edition", document.edition),
        ("Status", "Living document — the constitution of EdirasX"),
        ("Canonical source", "docs/edtechx/EDTECHX_EDITORIAL_BIBLE.md"),
        ("Authority", "Supreme. Where any document, design, or code conflicts with this Bible, the Bible wins until formally amended."),
        ("Amendment", "By entry in EDTECHX_DECISIONS.md, stating the principle changed, the reason, and the consequences."),
        ("Typography", f"Set in {SERIF} and {SANS}. Both are open-licence; Word will substitute if they are not installed."),
        ("Contents", "The table below is a live field. Press Ctrl+A then F9 (or right-click → Update Field) to populate page numbers."),
    ):
        cells = control.add_row().cells
        for index, text in enumerate((key, value)):
            cells[index].text = ""
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(text)
            run.font.size = Pt(9)
            run.font.color.rgb = INK_3 if index == 0 else INK
            _set_font(run._r, SANS)

    # Contents
    word.add_page_break()
    heading = _paragraph(word, space_after=10)
    run = heading.add_run("Contents")
    run.bold = True
    run.font.size = Pt(20)
    _set_font(run._r, SERIF)
    _outline_level(heading, 0)
    toc = _paragraph(word, space_after=6)
    _field(
        toc,
        ' TOC \\o "1-2" \\h \\z \\u ',
        "Right-click here and choose “Update Field” to build the contents.",
    )

    index = 1
    for chapter in document.front_matter:
        _chapter_heading(word, chapter, "", index)
        for block in chapter.blocks:
            _render_block(word, block)
        index += 1

    for chapter in document.chapters:
        part = next((p for p in document.parts if p.number == chapter.part), None)
        if part and not getattr(part, "_emitted", False):
            word.add_page_break()
            paragraph = _paragraph(word, space_before=120, space_after=6)
            run = paragraph.add_run(f"PART {part.number}")
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = ACCENT_2
            _set_font(run._r, SANS)
            title = _paragraph(word, space_after=8)
            run = title.add_run(part.title)
            run.bold = True
            run.font.size = Pt(28)
            _set_font(run._r, SERIF)
            _border(title, "bottom", "7A3FD6", size=12)
            standfirst = _paragraph(word, space_after=0)
            run = standfirst.add_run(typographic(part.standfirst))
            run.font.size = Pt(12)
            run.font.color.rgb = INK_2
            _set_font(run._r, SERIF)
            part._emitted = True
        _chapter_heading(word, chapter, f"Chapter {chapter.number}", index)
        for block in chapter.blocks:
            _render_block(word, block)
        index += 1

    for chapter in document.back_matter:
        _chapter_heading(word, chapter, "In closing", index)
        for block in chapter.blocks:
            _render_block(word, block)
        index += 1

    closing = _paragraph(word, space_before=18, space_after=0)
    _border(closing, "top", "D9D9D3", size=6)
    run = closing.add_run(
        "Generated from docs/edtechx/EDTECHX_EDITORIAL_BIBLE.md by tools/publisher. "
        "The Markdown source is canonical; this document is a publication artefact of it."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = INK_3
    _set_font(run._r, SANS)

    for part in document.parts:
        if hasattr(part, "_emitted"):
            delattr(part, "_emitted")

    word.save(output)
    return output


__all__ = ["render"]
